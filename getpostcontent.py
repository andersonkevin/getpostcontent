import requests
from bs4 import BeautifulSoup
from docx import Document
import re
import time

# Función para obtener el contenido del post desde la URL y estructurarlo, con manejo de errores
def get_post_content(url, retries=3):
    for i in range(retries):
        try:
            response = requests.get(url, timeout=10)  # Agregando un timeout
            response.raise_for_status()  # Levanta una excepción para códigos de estado HTTP 4xx/5xx
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Aquí identificamos el contenido principal del post.
            post_content = soup.find('article')  # Puede que necesites ajustar esto según la estructura de la página

            if post_content:
                return post_content
            else:
                print(f"No se pudo encontrar el contenido del post en la URL: {url}")
                return None
        except (requests.exceptions.ChunkedEncodingError, requests.exceptions.ConnectionError) as e:
            print(f"Error al procesar la URL {url}: {e}. Intentando de nuevo ({i+1}/{retries})...")
            time.sleep(5)  # Espera unos segundos antes de reintentar
        except requests.exceptions.RequestException as e:
            print(f"Error fatal con la URL {url}: {e}")
            return None
    print(f"Falló la obtención del contenido después de {retries} intentos.")
    return None

# Función para guardar el contenido en un archivo .docx manteniendo la estructura y agregando la URL
def save_to_docx(soup_content, url):
    doc = Document()
    
    # Añadir la URL al inicio del documento
    doc.add_paragraph(f'URL: {url}', style='Intense Quote')
    
    # Buscar el primer <h1> y usarlo como nombre de archivo
    h1_element = soup_content.find('h1')
    if h1_element:
        filename = re.sub(r'[^\w\s-]', '', h1_element.get_text().strip())  # Limpiar el nombre de archivo
        filename = filename.replace(' ', '_') + '.docx'

        # Añadir el H1 al documento como título
        doc.add_heading(h1_element.get_text(), level=1)
        
        # Eliminar contenido antes del H1 y añadir todo lo que sigue
        for element in h1_element.find_all_next(True):
            if element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.get_text(), level=4)
            elif element.name == 'h5':
                doc.add_heading(element.get_text(), level=5)
            elif element.name == 'h6':
                doc.add_heading(element.get_text(), level=6)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
            elif element.name == 'ul':  # Para listas no ordenadas
                for li in element.find_all('li'):
                    doc.add_paragraph(f'• {li.get_text()}', style='List Bullet')
            elif element.name == 'ol':  # Para listas ordenadas
                for li in element.find_all('li'):
                    doc.add_paragraph(f'{li.get_text()}', style='List Number')

        # Guardar el documento con el nombre basado en el H1
        doc.save(filename)
        print(f"El contenido ha sido guardado en {filename}")
    else:
        print("No se encontró un <h1> en el contenido.")

# Lista de URLs a procesar
urls = [
    # Añade aquí las URLs
   ]

# Procesar cada URL en la lista
for url in urls:
    print(f"Procesando URL: {url}")
    soup_content = get_post_content(url)
    if soup_content:
        save_to_docx(soup_content, url)
    else:
        print(f"No se pudo procesar la URL: {url}")
